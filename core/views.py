from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.conf import settings
from .models import Video
from .forms import VideoUploadForm
from .tasks import process_video, add_video_to_queue
from django.http import HttpResponse
from datetime import datetime, timedelta, time
import docx
import json
import logging
import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement, ns
from .models import Video
from itertools import groupby
from .tasks import process_video_task  # Importa a tarefa Celery
from transformers import AutoTokenizer, AutoModelForQuestionAnswering
import torch
from .tasks import video_queue

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')
log_message = lambda message: logging.info(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {message}")

# Inicialização do modelo e tokenizer
tokenizer = AutoTokenizer.from_pretrained("pierreguillou/bert-large-cased-squad-v1.1-portuguese")
model = AutoModelForQuestionAnswering.from_pretrained("pierreguillou/bert-large-cased-squad-v1.1-portuguese")

def gallery(request):
    videos = Video.objects.all()
    return render(request, 'core/gallery.html', {'videos': videos})

# def video_detail(request, video_id):
#     video = get_object_or_404(Video, id=video_id)
#     log_message(f"Obtendo os detalhes do vídeo {video.id}")
#         # Converte a string JSON em uma lista de dicionários, se necessário
#     if isinstance(video.transcription_phrases, str):
#         video.transcription_phrases = json.loads(video.transcription_phrases)

#     # Agrupa as palavras pelo orador, se a diarização estiver ativada e não houver erros
#     if video.diarize and video.word_timestamps and not video.error_on_diarization:
#         video.word_groups = group_words_by_speaker(video.word_timestamps)
#         log_message(f"Formados {len(video.word_groups)} grupos de palavras por orador.")
#     else:
#         log_message(f"Não foram formados grupos de palavras por orador.")
#         video.word_groups = None
    
#     # Converte o tempo de start e end para HH:MM:ss
#     for phrase in video.transcription_phrases:
#         phrase['start'] = format_time(float(phrase['start']))
#         phrase['end'] = format_time(float(phrase['end']))

#     # Renderiza o template passando o contexto com o vídeo
#     return render(request, 'core/video_detail.html', {'video': video})

def upload_video(request):
    if request.method == 'POST':
        form = VideoUploadForm(request.POST, request.FILES)
        if form.is_valid():
            video = form.save()
            messages.success(request, 'Upload bem-sucedido! O vídeo será processado em breve.')
            add_video_to_queue(video.id)
            return redirect('gallery')
    else:
        form = VideoUploadForm()
    return render(request, 'core/upload_video.html', {'form': form})

def export_to_docx(request, video_id):
    video = get_object_or_404(Video, id=video_id)
    
    doc = docx.Document()
    doc.add_heading(f'Transcrição do Vídeo: {video.file.name}', 0)

    if video.transcription:
        for line in video.transcription.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("Transcrição não disponível.")
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={video.file.name}.docx'
    doc.save(response)
    return response

def format_time(seconds):
    """Converte segundos para o formato HH:MM:ss"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = round(seconds % 60)
    return f"{hours:02}:{minutes:02}:{seconds:02}"

def video_detail_view(request, video_id):
    # Recupera o objeto Video do banco de dados
    video = get_object_or_404(Video, id=video_id)
    log_message(f"Obtendo os detalhes do vídeo {video.id}")

    if request.method == 'POST':
        question = request.POST.get('question')
        log_message(f"Pergunta recebida: {question}")
        
        if question:
            context = " ".join([phrase['text'] for phrase in video.transcription_phrases])
            log_message(f"Contexto gerado: {context[:500]}...")  # Mostra uma parte do contexto
            answer = answer_question(question, context)
            log_message(f"Resposta obtida: {answer}")
        else:
            answer = "Nenhuma pergunta foi fornecida."
            log_message("Nenhuma pergunta fornecida")
    else:
        answer = None
        log_message("Método GET utilizado, sem processamento de pergunta")
    # Converte o tempo de start e end para HH:MM:SS
    for phrase in video.transcription_phrases:
        phrase['start'] = format_time(float(phrase['start']))
        phrase['end'] = format_time(float(phrase['end']))

    # Renderiza o template passando o contexto com o vídeo
    return render(request, 'core/video_detail.html', {'video': video, 'answer': answer})

def format_duration(duration):
    """Converte um objeto `timedelta` ou `datetime.time` em uma string formatada HH:MM:SS"""
    if duration is None:
        return "00:00:00"  # Ou outra string padrão que você prefira para durações desconhecidas

    if isinstance(duration, timedelta):
        total_seconds = int(duration.total_seconds())
    elif isinstance(duration, time):
        total_seconds = duration.hour * 3600 + duration.minute * 60 + duration.second
    else:
        total_seconds = int(duration)
        
    hours, remainder = divmod(total_seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f'{hours:02}:{minutes:02}:{seconds:02}'

def convert_seconds_to_hms(seconds):
    """Converte tempo em segundos para o formato HH:MM:SS"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = int(seconds % 60)
    return f"{hours:02}:{minutes:02}:{seconds:02}"

def set_cell_vertical_alignment(cell, alignment):
    """Alinha verticalmente o conteúdo da célula ao centro"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(ns.qn('w:val'), alignment)
    tcPr.append(tcVAlign)

def set_cell_margins(cell, top=0, bottom=0):
    """Configura as margens internas superior e inferior da célula"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    tcMar.set(ns.qn('w:top'), str(top))
    tcMar.set(ns.qn('w:bottom'), str(bottom))
    tcPr.append(tcMar)

def generate_docx(request, video_id):
    video = get_object_or_404(Video, id=video_id)
    
    # Cria o documento
    doc = Document()

    # Configurações do documento - A4 é 210 x 297 mm
    section = doc.sections[0]
    section.page_width = Mm(210)   # 210 mm
    section.page_height = Mm(297)  # 297 mm
    section.left_margin = Mm(30)   # 30 mm para a margem esquerda
    section.right_margin = Mm(25)  # 25 mm para a margem direita
    section.top_margin = Mm(20)    # 25 mm para a margem superior
    section.bottom_margin = Mm(20) # 25 mm para a margem inferior

    # Título
    title = doc.add_heading(f'Transcrição do Vídeo: {video.file.name}', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Ajusta o tamanho da fonte do título
    run = title.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(14)

    # Cria a tabela para a transcrição
    if video.transcription_phrases:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True 
        table.allow_autofit = False
        table.width = Mm(section.page_width - section.left_margin - section.right_margin)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Início'
        for paragraph in hdr_cells[0].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[1].text = 'Texto'
        for paragraph in hdr_cells[1].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Define as células da primeira linha como cabeçalhos
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True

        # Repete a linha de cabeçalho em todas as páginas
        tbl = table._tbl
        tbl_header = tbl.xpath(".//w:tr")[0]
        tbl_header_pr = tbl_header.get_or_add_trPr()
        tbl_header_pr.append(OxmlElement('w:tblHeader'))

        for phrase in video.transcription_phrases:
            row_cells = table.add_row().cells
            # Converta o tempo de início para HH:MM:SS e remova espaços
            start_time = convert_seconds_to_hms(float(phrase['start'])).strip()
            row_cells[0].text = start_time
            row_cells[0].width = Mm(20)  # Ajusta a largura da coluna de início
            for paragraph in row_cells[0].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
            # Remove espaços e ajusta o alinhamento do texto
            text_content = str(phrase['text']).strip()
            row_cells[1].text = text_content
            row_cells[1].width = Mm(section.page_width - section.left_margin - section.right_margin - 20)
            for paragraph in row_cells[1].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Ajusta o alinhamento vertical e margens internas das células
            # for cell in row_cells:
                
        # Ajustar o tamanho da fonte da tabela
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        
                set_cell_vertical_alignment(cell, 'center')
                set_cell_margins(cell, top=Mm(2.5), bottom=Mm(2.5))  # 0,25 cm em pontos

    else:
        doc.add_paragraph("Transcrição não disponível.")

    # Gera o arquivo DOCX
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={video.file.name}.docx'
    doc.save(response)
    return response

def gallery_view(request):
    videos = Video.objects.all()
    for video in videos:
        # Certifique-se de que cada vídeo tem a duração formatada
        video.formatted_duration = format_duration(video.duration)
    return render(request, 'core/gallery.html', {'videos': videos})

def delete_transcription(request, video_id):
    video = get_object_or_404(Video, id=video_id)
    # Exclui o vídeo e todos os arquivos relacionados
    error = False
    try:
        video.file.delete()  # Exclui o arquivo de vídeo associado
    except:
        error = True
        log_message(f'Erro ao excluir o arquivo de vídeo: {video.file.name}')
    try:
        video.delete()       # Exclui o registro do vídeo no banco de dados
    except:
        error = True
        log_message(f'Erro ao excluir o registr do banco de dados: {video.file.name}')
        
    if not error:
        messages.success(request, f'Transcrição e arquivos do vídeo "{video.file.name}" excluídos com sucesso.')

    return redirect('gallery')  # Redireciona para a galeria após a exclusão

def group_words_by_speaker(diarization, word_timestamps):
    grouped = []
    current_speaker = None
    current_segment = {'speaker': None, 'start': None, 'end': None, 'text': ''}

    for segment in diarization:
        speaker = segment['speaker']
        start = segment['start']
        end = segment['end']

        # Filtra as palavras que estão dentro do intervalo de tempo do segmento de fala
        words_in_segment = [
            word for word in word_timestamps 
            if word['start'] >= start and word['end'] <= end
        ]

        # Se houver mudança de orador ou se for o primeiro segmento
        if speaker != current_speaker:
            # Salva o segmento atual (se não for o primeiro)
            if current_speaker is not None:
                grouped.append(current_segment)
            
            # Inicia um novo segmento
            current_speaker = speaker
            current_segment = {
                'speaker': speaker,
                'start': start,
                'end': end,
                'text': ' '.join([word['word'] for word in words_in_segment])
            }
        else:
            # Continua o segmento atual
            current_segment['end'] = end
            current_segment['text'] += ' ' + ' '.join([word['word'] for word in words_in_segment])

    # Adiciona o último segmento
    if current_segment['speaker'] is not None:
        grouped.append(current_segment)

    return grouped

def group_segments_by_speaker(diarization):
    grouped = []
    current_speaker = None
    current_segment = {'speaker': None, 'start': None, 'end': None}

    for segment in diarization:
        speaker = segment['speaker']
        start = segment['start']
        end = segment['end']

        # Se houver mudança de orador ou se for o primeiro segmento
        if speaker != current_speaker:
            # Salva o segmento atual (se não for o primeiro)
            if current_speaker is not None:
                grouped.append(current_segment)
            
            # Inicia um novo segmento
            current_speaker = speaker
            current_segment = {
                'speaker': speaker,
                'start': start,
                'end': end
            }
        else:
            # Continua o segmento atual
            current_segment['end'] = end

    # Adiciona o último segmento
    if current_segment['speaker'] is not None:
        grouped.append(current_segment)

    return grouped
def split_context(context, max_len=512):
    tokens = tokenizer.tokenize(context)
    return [" ".join(tokens[i:i+max_len]) for i in range(0, len(tokens), max_len)]

def answer_question(question, context):
    log_message("Iniciando processamento da pergunta")

    # Limitar o contexto ao máximo de 512 tokens (removendo tokens além do limite)
    inputs = tokenizer(question, context, truncation=True, max_length=512, add_special_tokens=True, return_tensors="pt")
    log_message(f"Inputs gerados: {inputs}")

    input_ids = inputs["input_ids"].tolist()[0]
    log_message(f"Input IDs: {input_ids}")

    outputs = model(**inputs)
    log_message(f"Outputs gerados: {outputs}")

    answer_start_scores = outputs.start_logits
    answer_end_scores = outputs.end_logits

    answer_start = torch.argmax(answer_start_scores)  # Índice de início da resposta
    answer_end = torch.argmax(answer_end_scores) + 1  # Índice de fim da resposta
    log_message(f"Answer start: {answer_start}, Answer end: {answer_end}")

    answer = tokenizer.convert_tokens_to_string(tokenizer.convert_ids_to_tokens(input_ids[answer_start:answer_end]))
    log_message(f"Resposta gerada: {answer}")

    return answer

def gallery_view(request):
    videos = Video.objects.all()
    video_queue_list = list(video_queue.queue)  # Converte a fila em uma lista para consulta
    for video in videos:
        # Verifica se o vídeo estava em processamento antes de um reinício e se não está na fila
        if video.in_process and video.id not in video_queue_list:
            video.queue_position = 1  # Considera como em transcrição, mesmo após o reinício
        elif video.id in video_queue_list:
            video.queue_position = video_queue_list.index(video.id) + 1  # Calcula a posição na fila
        else:
            video.queue_position = None  # Se o vídeo não estiver na fila, não tem posição
        # Certifique-se de que cada vídeo tem a duração formatada
        video.formatted_duration = format_duration(video.duration)
    return render(request, 'core/gallery.html', {'videos': videos})